"""Build a JIBEP-submission-ready .docx from the corrected Mizan markdown.

Journal formatting (JIBEP author guidelines):
  - Microsoft Word .docx
  - 11-point Times New Roman (serif)
  - 1.5 line spacing
  - APA 7th references (already in the source)
  - English (UK) (already in the source)
  - Tables/figures embedded in the text body, clearly labelled
"""
import re
import subprocess
import tempfile
import os
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = r"D:\01-claude\azim.cc\content\drafts\Mizan-JIBEP-Revised-June2026.md"
OUT = r"D:\01-claude\azim.cc\content\drafts\Mizan-JIBEP-Revised-June2026.docx"

BODY_PT = 11
TABLE_PT = 10
HEADING_PT = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 12,
              "Heading 4": 11, "Heading 5": 11, "Heading 6": 11}

# 1. Preprocess markdown: drop standalone '---' thematic breaks (section separators)
with open(SRC, encoding="utf-8") as f:
    lines = f.readlines()
kept = [ln for ln in lines if ln.strip() != "---"]
md = "".join(kept)

with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as tf:
    tf.write(md)
    tmp_md = tf.name

# 2. pandoc -> docx
try:
    subprocess.run(
        ["pandoc", tmp_md, "-f", "gfm", "-t", "docx", "-o", OUT],
        check=True, capture_output=True, text=True,
    )
finally:
    os.unlink(tmp_md)

# 3. Post-process with python-docx
doc = Document(OUT)

def set_run_font(run, size=None, name="Times New Roman"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)

# Base 'Normal' style
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(BODY_PT)
normal.paragraph_format.line_spacing = 1.5

# Body paragraphs (title + author block centred until first section heading)
seen_first_heading2 = False
in_references = False
for para in doc.paragraphs:
    sname = para.style.name if para.style else "Normal"
    if sname == "Heading 2":
        seen_first_heading2 = True
        in_references = para.text.strip().lower() == "references"

    if sname.startswith("Heading"):
        for run in para.runs:
            set_run_font(run, HEADING_PT.get(sname, BODY_PT))
            run.font.bold = True
        if sname == "Heading 1":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif "Source Code" in sname or "Verbatim" in sname:
        # ASCII schematic: keep monospaced so alignment is preserved
        para.paragraph_format.line_spacing = 1.0
        for run in para.runs:
            set_run_font(run, 9, name="Consolas")
    else:
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            set_run_font(run, BODY_PT)
        if not seen_first_heading2 and para.text.strip():
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if in_references and para.text.strip():
            # APA 7th hanging indent for reference entries
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)

# Tables: 10pt TNR, single spacing for compact fit
for table in doc.tables:
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    set_run_font(run, TABLE_PT)

# Page: A4, 1-inch margins (UK publisher)
for section in doc.sections:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

doc.save(OUT)
print("Saved:", OUT)
print("Tables:", len(doc.tables), "| Paragraphs:", len(doc.paragraphs))
